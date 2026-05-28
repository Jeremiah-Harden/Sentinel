"""
build_resumes_bulk.py — Generate one tailored .docx resume per job in the Excel.

Output: resumes/{company_slug}_resume.docx  (or _2, _3 for duplicates)

How tailoring works (keyword-scoring, runs fully offline — no API calls):
  1. classify() counts how many cyber / swe / data keywords appear in the
     job title + description + tags combined.
  2. The category with the highest count wins. If all are zero → "general".
  3. Each category maps to a different summary paragraph and a different
     project order so the most relevant project always leads.

Why keyword scoring instead of AI classification?
  Speed and cost. We might generate 100+ resumes in one run. An LLM API call
  per job would be slow and expensive. Keyword counting is deterministic,
  instant, and accurate enough for the four broad categories we care about.

Why a fresh Document() per resume instead of copying a template?
  Templates require keeping a .docx binary in sync with the code. Generating
  from code means the layout is fully version-controlled and reproducible —
  re-run the script at any time to regenerate all resumes with updated content.

Duplicate company handling:
  Multiple jobs at the same company get _resume.docx, _resume_2.docx, etc.
  so no file is silently overwritten.
"""

import re
import shutil
import sys
import tempfile
from pathlib import Path

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Summaries ─────────────────────────────────────────────────────────────────

SUMMARIES = {
    "cyber": (
        "Cybersecurity student at Kennesaw State University pursuing a Bachelor's degree, with hands-on "
        "experience in threat detection, log analysis, and incident investigation through self-directed "
        "security labs. Proficient in Python, Kali Linux, and security tools including Wireshark, Nmap, "
        "and Splunk. Adept at building monitoring environments and applying analytical thinking to "
        "real-world cyber threat scenarios."
    ),
    "swe": (
        "Cybersecurity student at Kennesaw State University pursuing a Bachelor's degree, with hands-on "
        "experience in Python scripting, REST API integration, and automated data pipeline engineering. "
        "Proficient in building CLI tools and agentic workflow systems using the WAT framework "
        "(Workflows, Agents, Tools). Skilled in Linux, network programming, and applying software "
        "engineering practices to security and automation challenges."
    ),
    "data": (
        "Cybersecurity student at Kennesaw State University pursuing a Bachelor's degree, with hands-on "
        "experience in Python scripting, structured data extraction, and workflow automation. Proficient "
        "in building agentic pipelines using the WAT framework (Workflows, Agents, Tools) and the "
        "FireCrawl API. Skilled in data analysis, security monitoring, and applying technical "
        "problem-solving to engineering challenges."
    ),
    "general": (
        "Cybersecurity student at Kennesaw State University pursuing a Bachelor's degree, with hands-on "
        "experience in threat detection, Python automation, and agentic workflow engineering. Proficient "
        "in Kali Linux, Wireshark, Nmap, and Splunk. Experienced in building security labs and "
        "automation pipelines, with a strong foundation in applying analytical thinking to real-world "
        "technical challenges."
    ),
}

# ── Projects ──────────────────────────────────────────────────────────────────

PROJECTS = {
    "honeypot": {
        "title": "Raspberry Pi Honeypot and Threat Monitoring Lab",
        "bullets": [
            "Built a honeypot environment using a Raspberry Pi 5 and mini PC to simulate vulnerable ports and network services.",
            "Captured and analyzed login attempts and automated network scanning activity.",
            "Investigated attack behavior using Linux logs, Nmap, and TCPdump.",
        ],
    },
    "portscanner": {
        "title": "Network Port Scanner (Python and NMap)",
        "bullets": [
            "Developed a network port scanner using Python and the Nmap library to identify open ports and assess security vulnerabilities on targeted IP addresses for use in penetration testing.",
        ],
    },
    "malware": {
        "title": "Malware Analysis Lab",
        "bullets": [
            "Designed a virtualized malware analysis environment using VMware with isolated Windows/Linux virtual machines.",
            "Conducted dynamic malware investigations and network traffic analysis using Wireshark.",
        ],
    },
    "agentive": {
        "title": "Agentive Workflow Automation (WAT Framework)",
        "bullets": [
            "Engineered a multi-phase data pipeline using the WAT framework (Workflows, Agents, Tools), with Claude AI as the probabilistic orchestration layer and Python CLI scripts as the deterministic execution layer.",
            "Automated structured web data extraction across 20 search terms via the FireCrawl API, collecting and deduplicating 145 unique records into a formatted Excel deliverable.",
            "Designed the tool layer to be modular and reusable across any query or target platform, demonstrating applied knowledge of AI-driven automation and REST API integration.",
        ],
    },
}

PROJECT_ORDER = {
    "cyber":   ["honeypot", "portscanner", "malware", "agentive"],
    "swe":     ["agentive", "portscanner", "honeypot", "malware"],
    "data":    ["agentive", "portscanner", "honeypot", "malware"],
    "general": ["honeypot", "portscanner", "malware", "agentive"],
}

SKILLS = [
    ("Python",    "Kali Linux",            "Linux, Windows, MacOS"),
    ("Basic SQL", "Splunk",                "TCP/IP, Port Scanning"),
    ("Java",      "SQLmap · FireCrawl",    "Network Troubleshooting"),
    ("C#",        "Wireshark and NMap",    "Security Monitoring"),
]

# Keyword sets for classify(). Each set is checked against the combined
# job title + description + tags text. The category whose keywords appear
# most often wins. Sets are intentionally broad — a false positive to "swe"
# is better than falling through to "general" and losing a tailored summary.
CYBER_KW = {
    "security", "cyber", "threat", "incident", "soc", "vulnerability",
    "penetration", "malware", "forensic", "siem", "splunk", "nmap",
    "wireshark", "firewall", "risk", "compliance", "audit", "defense",
    "intrusion", "detection", "blue team", "red team",
}
SWE_KW = {
    "software", "developer", "engineer", "python", "java", "javascript",
    "react", "backend", "frontend", "fullstack", "api", "database", "sql",
    "devops", "cloud", "automation", "pipeline", "web", "mobile", "node",
    "typescript", "golang", "ruby", "rails", "django", "flask",
}
DATA_KW = {
    "data", "analytics", "machine learning", "ml", "ai", "artificial",
    "intelligence", "etl", "warehouse", "pandas", "numpy", "spark",
    "tableau", "power bi", "extraction", "scraping", "pipeline",
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def classify(title, description, tags):
    """Score a job against three keyword sets and return the winning category.

    Ties between cyber and swe go to cyber — security tailoring is always
    more specific and valuable for Jeremiah's profile. Data beats swe only
    when it scores strictly higher. Cyber always wins a tie with data too.
    Falls through to "general" when no keywords match at all.
    """
    text = f"{title} {description} {tags}".lower()
    cy = sum(1 for kw in CYBER_KW if kw in text)
    sw = sum(1 for kw in SWE_KW if kw in text)
    da = sum(1 for kw in DATA_KW if kw in text)
    if cy >= sw and cy >= da:
        return "cyber"
    if da > sw:
        return "data"
    if sw > 0:
        return "swe"
    return "general"


def slugify(text):
    """Convert a company name to a safe filename fragment.

    Strips punctuation, collapses whitespace to underscores, and caps at 40
    characters so filenames stay readable. The 40-char cap also avoids hitting
    Windows MAX_PATH issues when the output directory has a long path.
    """
    text = (text or "unknown").lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_-]+", "_", text)
    return text[:40].strip("_")


def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBdr.append(el)
    tcPr.append(tcBdr)


def add_bottom_rule(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    add_bottom_rule(p)


def cell_section_header(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    add_bottom_rule(p)


def body_para(doc, text, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.add_run(text).font.size = Pt(10)


def proj_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.add_run(text).font.size = Pt(10)


def job_line(doc, company, date):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(company)
    r1.bold = True
    r1.underline = True
    r1.font.size = Pt(10)
    r2 = p.add_run("\t" + date)
    r2.bold = True
    r2.font.size = Pt(10)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")
    tabs.append(tab)
    pPr.append(tabs)


# ── Resume builder ────────────────────────────────────────────────────────────

def build_resume(output_path, summary, project_order):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Name
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Jeremiah Harden")
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.add_run("404-840-3083 · Jeremiah.Harden11@gmail.com · Kennesaw, GA · Kennesaw State").font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.add_run("LinkedIn: linkedin.com/in/jeremiah-harden-50ba4331a").font.size = Pt(10)

    # Summary
    section_header(doc, "SUMMARY")
    body_para(doc, summary, space_before=2, space_after=2)

    # Projects
    section_header(doc, "PROJECTS")
    for key in project_order:
        proj = PROJECTS[key]
        proj_heading(doc, proj["title"])
        for b in proj["bullets"]:
            add_bullet(doc, b)

    # Skills
    section_header(doc, "SKILLS AND PROGRAMMING LANGUAGES")
    tbl = doc.add_table(rows=len(SKILLS), cols=3)
    tbl.style = "Table Grid"
    for r_idx, row_data in enumerate(SKILLS):
        for c_idx, text in enumerate(row_data):
            cell = tbl.rows[r_idx].cells[c_idx]
            remove_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.add_run("• " + text).font.size = Pt(10)

    # Work experience
    section_header(doc, "WORK EXPERIENCE")
    job_line(doc, "Executive Parking Services - Valet", "September 2025 - Present")
    body_para(doc,
        "As a Valet at Executive Parking Services, I managed high-value vehicles in fast-paced environments, "
        "such as Capital Grille, requiring strong attention to detail, accountability, and risk awareness. "
        "I coordinated vehicle tracking and retrieval using Oobeo while maintaining professionalism under "
        "time-sensitive conditions. This role strengthened my reliability, communication skills, and ability "
        "to perform in high-responsibility situations.",
        space_before=2, space_after=4)

    job_line(doc, "The Athletes Foot Corporate Office - Retail Associate", "January 2025 - August 2025")
    body_para(doc,
        "At The Athletes Foot Corporate Store, I gained exposure to data-driven marketing strategies and "
        "learned how statistical trends influence product promotion and customer targeting. I contributed "
        "to event planning and execution, supporting promotional campaigns and product launches in "
        "professional corporate environments.",
        space_before=2, space_after=4)

    # Education & Extracurricular (two-column)
    bottom = doc.add_table(rows=1, cols=2)
    bottom.style = "Table Grid"
    left = bottom.rows[0].cells[0]
    right = bottom.rows[0].cells[1]

    left.paragraphs[0].clear()
    cell_section_header(left, "EDUCATION & CERTIFICATIONS")

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Kennesaw State University")
    r.bold = True; r.underline = True; r.font.size = Pt(10)

    for line in ["Major: Cybersecurity", "Graduation: Summer 2027"]:
        p = left.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.add_run("• " + line).font.size = Pt(10)

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run("• Relevant Coursework: ")
    r1.font.size = Pt(10); r1.bold = True
    p.add_run(
        "Programming Problem Solving — developed Python scripts to parse and organize system logs, "
        "identifying and visualizing security incidents for Blackthorne's Internship Challenge"
    ).font.size = Pt(10)

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Shiloh High School")
    r.bold = True; r.underline = True; r.font.size = Pt(10)

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.add_run("• High School Diploma").font.size = Pt(10)

    right.paragraphs[0].clear()
    cell_section_header(right, "EXTRACURRICULAR ACTIVITIES")

    p = right.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Black Men in Cyber")
    r.bold = True; r.underline = True; r.font.size = Pt(10)

    for line in [
        "Active member of a collaborative community of cybersecurity-focused students sharing tools "
        "including Nmap, Shodan, Censys, VirusTotal, and more.",
        "Participated in labs and hands-on projects facilitated by TryHackMe and HackTheBox.",
    ]:
        p = right.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("• " + line).font.size = Pt(10)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    excel_path = ".tmp/tech_internships.xlsx"
    output_dir = Path("resumes")

    print("=" * 65)
    print("  BULK RESUME GENERATOR")
    print(f"  Source: {excel_path}")
    print(f"  Output: {output_dir}/")
    print("=" * 65)

    # Copy to a temp file before opening. On Windows, openpyxl can't open a
    # .xlsx that Excel currently has locked. The copy bypasses the file lock
    # so the script works even when the spreadsheet is open on your desktop.
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp_path = tmp.name
    shutil.copy2(excel_path, tmp_path)
    wb = load_workbook(tmp_path, read_only=True, data_only=True)
    ws = wb.active
    headers = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    print(f"\n  {len(rows)} jobs loaded from Excel\n")

    used_slugs: dict[str, int] = {}
    success = 0
    skipped = 0

    for i, row in enumerate(rows, start=1):
        data = dict(zip(headers, row))
        company   = str(data.get("Company Name") or "Unknown").strip()
        job_title = str(data.get("Job Title") or "Unknown").strip()
        desc      = str(data.get("Job Description") or "")
        reqs      = str(data.get("Requirements") or "")
        tags      = str(data.get("Tags / Skills") or "")

        job_type = classify(job_title, desc + " " + reqs, tags)
        summary  = SUMMARIES[job_type]
        order    = PROJECT_ORDER[job_type]

        slug = slugify(company)
        if slug in used_slugs:
            used_slugs[slug] += 1
            filename = f"{slug}_resume_{used_slugs[slug]}.docx"
        else:
            used_slugs[slug] = 1
            filename = f"{slug}_resume.docx"

        out = output_dir / filename

        try:
            build_resume(str(out), summary, order)
            label = f"[{job_type.upper():7s}]"
            print(f"  [{i:03d}/{len(rows)}] {label}  {company[:30]:<30}  →  {filename}")
            success += 1
        except Exception as exc:
            print(f"  [{i:03d}/{len(rows)}] ERROR  {company[:30]:<30}  ✗  {exc}")
            skipped += 1

    print()
    print("=" * 65)
    print(f"  DONE  |  {success} resumes saved  |  {skipped} errors")
    print(f"  Location: {output_dir.resolve()}")
    print("=" * 65)


if __name__ == "__main__":
    main()
