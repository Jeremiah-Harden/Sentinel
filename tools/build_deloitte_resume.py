"""
build_deloitte_resume.py — Generate a Deloitte-tailored resume as a .docx file.

Why python-docx instead of a Word template?
  A template requires editing the .docx binary manually and breaks whenever
  you change the structure. Generating from code means the layout, fonts,
  and content are all in one place and reproducible — re-run the script to
  regenerate a perfectly formatted resume at any time.

Why direct OOXML manipulation (OxmlElement, qn)?
  python-docx's high-level API doesn't expose everything. Things like removing
  table cell borders and adding paragraph border rules require dropping down to
  the underlying Office Open XML (OOXML) format. This is verbose but precise —
  we're setting the same XML attributes Word would write if you clicked
  "Format → Borders and Shading" in the UI.

Output: resumes/deloitte_resume.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def remove_cell_borders(cell):
    """Remove all borders from a table cell using raw OOXML.

    python-docx doesn't expose cell border removal through its public API,
    so we manipulate the underlying XML directly. The `qn()` helper converts
    Python-friendly attribute names to the namespaced XML form Word expects
    (e.g., 'w:val' → '{http://schemas.openxmlformats.org/...}val').

    Used for the skills table and the two-column education/extracurricular table
    so they look like normal content rather than bordered table cells.
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBdr.append(el)
    tcPr.append(tcBdr)


def add_bottom_rule(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    add_bottom_rule(p)
    return p


def cell_section_header(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    add_bottom_rule(p)
    return p


def body(doc, text, size=10, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def proj_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    return p


def bullet(container, text, size=10):
    p = container.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def job_line(doc, company, date, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(company)
    r1.bold = True
    r1.underline = True
    r1.font.size = Pt(size)
    r2 = p.add_run('\t' + date)
    r2.bold = True
    r2.font.size = Pt(size)
    # Right-align the tab stop
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9360')
    tabs.append(tab)
    pPr.append(tabs)


# ── Document ──────────────────────────────────────────────────────────────────

def build(output_path: str):
    doc = Document()

    # Margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── NAME & CONTACT ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('Jeremiah Harden')
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.add_run('404-840-3083 · Jeremiah.Harden11@gmail.com · Kennesaw, GA · Kennesaw State').font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.add_run('LinkedIn: linkedin.com/in/jeremiah-harden-50ba4331a').font.size = Pt(10)

    # ── SUMMARY ───────────────────────────────────────────────────────────
    section_header(doc, 'SUMMARY')
    body(doc,
        'Cybersecurity student at Kennesaw State University pursuing a Bachelor\'s degree, with hands-on '
        'experience in threat detection, log analysis, and incident investigation through self-directed security '
        'labs and academic challenges. Proficient in Python, Kali Linux, and security tools including Wireshark, '
        'Nmap, and Splunk. Experienced in agentic workflow automation and data pipeline engineering using '
        'AI-driven orchestration frameworks.',
        space_before=2, space_after=2)

    # ── PROJECTS ──────────────────────────────────────────────────────────
    section_header(doc, 'PROJECTS')

    proj_title(doc, 'Raspberry Pi Honeypot and Threat Monitoring Lab')
    bullet(doc, 'Built a honeypot environment using a Raspberry Pi 5 and mini PC to simulate vulnerable ports and network services.')
    bullet(doc, 'Captured and analyzed login attempts and automated network scanning activity.')
    bullet(doc, 'Investigated attack behavior using Linux logs, Nmap, and TCPdump.')

    proj_title(doc, 'Network Port Scanner (Python and NMap)')
    bullet(doc, 'Developed a network port scanner using Python and the Nmap library to identify open ports and assess security vulnerabilities on targeted IP addresses for use in penetration testing.')

    proj_title(doc, 'Malware Analysis Lab')
    bullet(doc, 'Designed a virtualized malware analysis environment using VMware with isolated Windows/Linux virtual machines.')
    bullet(doc, 'Conducted dynamic malware investigations and network traffic analysis using Wireshark.')

    proj_title(doc, 'Agentive Workflow Automation (WAT Framework)')
    bullet(doc, 'Engineered a multi-phase data pipeline using the WAT framework (Workflows, Agents, Tools), with Claude AI as the probabilistic orchestration layer and Python CLI scripts as the deterministic execution layer.')
    bullet(doc, 'Automated structured web data extraction across 20 search terms via the FireCrawl API, collecting and deduplicating 145 unique records into a formatted Excel deliverable.')
    bullet(doc, 'Designed the tool layer to be modular and reusable across any query or target platform, demonstrating applied knowledge of AI-driven automation and REST API integration.')

    # ── SKILLS ────────────────────────────────────────────────────────────
    section_header(doc, 'SKILLS AND PROGRAMMING LANGUAGES')

    skills = [
        ('Python',    'Kali Linux',              'Linux, Windows, MacOS'),
        ('Basic SQL', 'Splunk',                  'TCP/IP, Port Scanning'),
        ('Java',      'SQLmap · FireCrawl',       'Network Troubleshooting'),
        ('C#',        'Wireshark and NMap',       'Security Monitoring'),
    ]

    tbl = doc.add_table(rows=len(skills), cols=3)
    tbl.style = 'Table Grid'
    for r_idx, row_data in enumerate(skills):
        for c_idx, text in enumerate(row_data):
            cell = tbl.rows[r_idx].cells[c_idx]
            remove_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run('• ' + text)
            run.font.size = Pt(10)

    # ── WORK EXPERIENCE ───────────────────────────────────────────────────
    section_header(doc, 'WORK EXPERIENCE')

    job_line(doc, 'Executive Parking Services - Valet', 'September 2025 - Present')
    body(doc,
        'As a Valet at Executive Parking Services, I managed high-value vehicles in fast-paced environments, '
        'such as Capital Grille, requiring strong attention to detail, accountability, and risk awareness. '
        'I coordinated vehicle tracking and retrieval using Oobeo while maintaining professionalism under '
        'time-sensitive conditions. This role strengthened my reliability, communication skills, and ability '
        'to perform in high-responsibility situations.',
        space_before=2, space_after=4)

    job_line(doc, 'The Athletes Foot Corporate Office - Retail Associate', 'January 2025 - August 2025')
    body(doc,
        'At The Athletes Foot Corporate Store, I gained exposure to data-driven marketing strategies and learned '
        'how statistical trends influence product promotion and customer targeting. I contributed to event planning '
        'and execution, supporting promotional campaigns and product launches in professional corporate environments.',
        space_before=2, space_after=4)

    # ── EDUCATION & EXTRACURRICULAR (two-column) ──────────────────────────
    bottom = doc.add_table(rows=1, cols=2)
    bottom.style = 'Table Grid'

    left = bottom.rows[0].cells[0]
    right = bottom.rows[0].cells[1]

    # Left: Education
    left.paragraphs[0].clear()
    cell_section_header(left, 'EDUCATION & CERTIFICATIONS')

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('Kennesaw State University')
    r.bold = True
    r.underline = True
    r.font.size = Pt(10)

    for line in ['Major: Cybersecurity', 'Graduation: Summer 2027']:
        p = left.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.add_run('• ' + line).font.size = Pt(10)

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run('• Relevant Coursework: ')
    r1.font.size = Pt(10)
    r1.bold = True
    r2 = p.add_run(
        'Programming Problem Solving — developed Python scripts to parse and organize system logs, '
        'identifying and visualizing security incidents for Blackthorne\'s Internship Challenge'
    )
    r2.font.size = Pt(10)

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('Shiloh High School')
    r.bold = True
    r.underline = True
    r.font.size = Pt(10)

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.add_run('• High School Diploma').font.size = Pt(10)

    # Right: Extracurricular
    right.paragraphs[0].clear()
    cell_section_header(right, 'EXTRACURRICULAR ACTIVITIES')

    p = right.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('Black Men in Cyber')
    r.bold = True
    r.underline = True
    r.font.size = Pt(10)

    for line in [
        'Active member of a collaborative community of cybersecurity-focused students sharing tools including '
        'Nmap, Shodan, Censys, VirusTotal, and more.',
        'Participated in labs and hands-on projects facilitated by TryHackMe and HackTheBox.',
    ]:
        p = right.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.add_run('• ' + line).font.size = Pt(10)

    # Save
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f'Saved: {out}')


if __name__ == '__main__':
    build('resumes/deloitte_resume.docx')
